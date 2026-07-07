from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)

def bullet(text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        before = text[:idx]
        after = text[idx+len(bold_part):]
        if before:
            p.add_run(before).font.size = Pt(11)
        r = p.add_run(bold_part)
        r.bold = True
        r.font.size = Pt(11)
        if after:
            p.add_run(after).font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    p.paragraph_format.left_indent = Inches(0.3)
    shading = p._element
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = shading.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)

def space():
    doc.add_paragraph()

# ─── 타이틀 ───
title = doc.add_heading('PostFab 임베딩 파인튜닝 계획서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph('PostFab Multi-Agent RAG 검색 정확도 향상 프로젝트')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph('작성일: 2026년 6월').alignment == WD_ALIGN_PARAGRAPH.CENTER
space()

# ─── 1. 목표 ───
h1('1. 목표')
body('BAAI/bge-m3 임베딩 모델을 후공정(Post-Fab) 도메인 데이터로 파인튜닝하여 RAG 검색 정확도를 수치로 증명한다.')
space()

# ─── 2. 사용 기술 ───
h1('2. 사용 기술')

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '항목'
hdr[1].text = '내용'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(11)

rows = [
    ('임베딩 모델', 'BAAI/bge-m3 (다국어, 한국어 강함)'),
    ('파인튜닝 라이브러리', 'sentence-transformers 3.4.1'),
    ('손실 함수', 'MultipleNegativesRankingLoss'),
    ('평가 도구', 'InformationRetrievalEvaluator'),
    ('평가 지표', 'Accuracy@3, NDCG@10 (Before/After 비교)'),
    ('학습 환경', 'Google Colab Pro (T4 GPU) — 예상 소요 5~15분'),
    ('Q&A 자동 생성', 'Claude API (용어당 3~5개 생성)'),
]
for k, v in rows:
    row = table.add_row().cells
    row[0].text = k
    row[1].text = v
    for cell in row:
        cell.paragraphs[0].runs[0].font.size = Pt(11)

space()

# ─── 3. 전체 흐름 ───
h1('3. 전체 흐름')
code('데이터 준비 → Q&A 자동 생성 → Train/Test 분리 → 파인튜닝 → Before/After 평가 → 시스템 연결')
space()

# ─── 4. 데이터 준비 ───
h1('4. 데이터 준비')

h2('4-1. 데이터 소스 (4가지)')
bullet('SK하이닉스 뉴스룸 크롤링', 'SK하이닉스 뉴스룸')
bullet('Advantest 용어집 (Claude가 한국어 번역)', 'Advantest 용어집')
bullet('Claude API가 생성한 후공정 용어 설명', 'Claude API')
bullet('현업 용어 (짜투리랏, R/S 등)', '현업 용어')

space()
h2('4-2. Q&A 자동 생성 방식')
body('각 용어 설명을 Claude API에 입력하여 현장 엔지니어가 실제로 물어볼 법한 질문 3~5개를 자동 생성한다.')
space()
body('예시:')
code('입력: "짜투리랏 = 박스 구성 후 남은 랏"\n\n생성:\n  Q: "짜투리랏이 뭐야?"\n  Q: "박스 구성하고 남은 랏은 어떻게 불러?"\n  Q: "잔여랏이랑 짜투리랏은 같은 말이야?"')

space()
h2('4-3. Train / Test 분리')

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
h = table2.rows[0].cells
for i, t in enumerate(['데이터 종류', 'Train', 'Test']):
    h[i].text = t
    h[i].paragraphs[0].runs[0].bold = True
    h[i].paragraphs[0].runs[0].font.size = Pt(11)

rows2 = [
    ('공개 용어 (FDC, Yield, Burn-in 등)', '80%', '20%'),
    ('현업 용어 (짜투리랏 등)', 'Train만 사용 또는 유사어 쌍으로 분리', '-'),
]
for r in rows2:
    row = table2.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

space()

# ─── 5. 파인튜닝 ───
h1('5. 파인튜닝')

h2('5-1. MultipleNegativesRankingLoss 원리')
body('(query, positive_doc) 쌍만 준비하면 된다. 배치 내 다른 문서들이 자동으로 네거티브 역할을 하므로 별도의 네거티브 샘플 구성 불필요.')
space()
code('배치 예시 (batch_size=4):\n\n  (질문1, 문서1) ← 포지티브\n  (질문2, 문서2) ← 포지티브\n  (질문3, 문서3) ← 포지티브\n  (질문4, 문서4) ← 포지티브\n\n  → 질문1 기준: 문서2,3,4는 자동으로 네거티브')

space()
h2('5-2. 핵심 코드')
code("""from sentence_transformers import SentenceTransformer, losses, InputExample
from torch.utils.data import DataLoader

model = SentenceTransformer("BAAI/bge-m3")

train_examples = [
    InputExample(texts=["짜투리랏이 뭐야?", "짜투리랏은 박스 구성 후 남은 랏..."]),
    InputExample(texts=["FDC CRITICAL 대응법?", "즉시 LOT 홀드 처리..."]),
    # ...
]

loader = DataLoader(train_examples, batch_size=16, shuffle=True)
loss = losses.MultipleNegativesRankingLoss(model)

model.fit(
    train_objectives=[(loader, loss)],
    epochs=3,
    output_path="./postfab_embedding_model"
)""")
space()

# ─── 6. 평가 ───
h1('6. 평가 (Before / After)')

h2('6-1. 평가 구조')
body('InformationRetrievalEvaluator는 세 가지 딕셔너리를 받아 자동으로 채점한다.')
space()
code("""queries       = {"q0": "짜투리랏이 뭐야?", "q1": "FDC CRITICAL 대응법?", ...}
corpus        = {"d0": "짜투리랏은 박스 구성 후...", "d1": "즉시 LOT 홀드...", ...}
relevant_docs = {"q0": {"d0"}, "q1": {"d1"}, ...}  # 정답 매핑""")

space()
h2('6-2. 평가 지표 설명')

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
for i, t in enumerate(['지표', '의미', '포트폴리오 활용']):
    table3.rows[0].cells[i].text = t
    table3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    table3.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)

rows3 = [
    ('Accuracy@3', '상위 3개 안에 정답 문서가 있는 비율.\n100번 질문 시 몇 번 정답 찾냐.', '주요 지표 — 직관적, 임원도 이해 가능'),
    ('NDCG@10', '정답이 1위에 가까울수록 높은 점수.\n순위까지 고려하는 정교한 지표.', '보조 지표 — 기술적 깊이 증명'),
]
for r in rows3:
    row = table3.add_row().cells
    for i, val in enumerate(r):
        row[i].text = val
        row[i].paragraphs[0].runs[0].font.size = Pt(10)

space()
h2('6-3. 목표 수치 (예시)')
code('파인튜닝 전: Accuracy@3 = 0.52\n파인튜닝 후: Accuracy@3 = 0.78\n→ 약 26%p 향상 → 포트폴리오 핵심 수치')
space()

# ─── 7. 시스템 연결 ───
h1('7. 시스템 연결')
body('파인튜닝 완료된 모델을 기존 RAG 파이프라인에 교체하면 전체 시스템에 즉시 적용된다.')
space()
code('# src/rag/build_vectorstore.py\n\n# 기존\nmodel = SentenceTransformer("BAAI/bge-m3")\n\n# 교체\nmodel = SentenceTransformer("./postfab_embedding_model")')
space()

# ─── 8. 작업 환경 ───
h1('8. 작업 환경 분리')

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
for i, t in enumerate(['로컬 (VSCode)', 'Google Colab Pro']):
    table4.rows[0].cells[i].text = t
    table4.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    table4.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)

row = table4.add_row().cells
row[0].text = '데이터 준비 및 전처리\nQ&A 생성 코드 작성\n시스템 최종 연결 및 테스트'
row[1].text = '학습 코드 (.ipynb) 실행\nT4 GPU로 파인튜닝 (5~15분)\n학습된 모델 다운로드'
for cell in row:
    cell.paragraphs[0].runs[0].font.size = Pt(10)

space()

# ─── 9. 실행 순서 ───
h1('9. 실행 순서 (체크리스트)')

steps = [
    '용어 데이터 수집 (뉴스룸 링크 + 현업 용어)',
    'Claude API로 Q&A 자동 생성 스크립트 작성',
    'Train(80%) / Test(20%) 분리',
    'Colab Pro에서 BAAI/bge-m3 파인튜닝 실행',
    'Before/After Accuracy@3, NDCG@10 수치 확인',
    '모델 다운로드 → src/rag/build_vectorstore.py 교체',
    '전체 시스템 동작 확인 및 데모 준비',
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s).font.size = Pt(11)

space()

doc.save('PostFab_임베딩파인튜닝_계획서.docx')
print("완료!")
